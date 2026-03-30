from fastapi import FastAPI, UploadFile, File, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from typing import List, Optional
import pdfplumber
import json
import uuid
import os

from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer,
    ListFlowable, ListItem
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import A4

from groq_client import analyze_legal_text
import auth as _auth
from email_service import send_reset_email, send_email, SMTP_HOST, SMTP_PORT, SMTP_USER, SMTP_PASS

# ── Supabase client (shared from auth module) ─────────────────────────────────
from supabase import create_client as _sb_create
_SUPABASE_URL = os.getenv("SUPABASE_URL", "")
_SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_KEY", "")
_sb = _sb_create(_SUPABASE_URL, _SUPABASE_KEY) if _SUPABASE_URL and _SUPABASE_KEY else None
import io
from docx import Document as DocxDocument

app = FastAPI()

# ─────────────────────────────────────────
# STORE DOCUMENT FOR CHAT  (unchanged)
# ─────────────────────────────────────────
DOCUMENT_CONTEXT = []

# ─────────────────────────────────────────
# CORS  (unchanged)
# ─────────────────────────────────────────
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─────────────────────────────────────────
# JWT GUARD
# ─────────────────────────────────────────
_bearer = HTTPBearer(auto_error=False)

def require_user(creds: Optional[HTTPAuthorizationCredentials] = Depends(_bearer)):
    if not creds:
        raise HTTPException(status_code=401, detail="Not authenticated. Please sign in.")
    user = _auth.get_user_from_token(creds.credentials)
    if not user:
        raise HTTPException(status_code=401, detail="Session expired. Please sign in again.")
    return user

# ─────────────────────────────────────────
# REQUEST SCHEMAS
# ─────────────────────────────────────────
class RegisterBody(BaseModel):
    name: str
    email: str
    password: str

class LoginBody(BaseModel):
    email: str
    password: str

class ChangePasswordBody(BaseModel):
    old_password: str
    new_password: str

class UpdateProfileBody(BaseModel):
    name: str

class ForgotPasswordBody(BaseModel):
    email: str

class ResetPasswordBody(BaseModel):
    token: str
    new_password: str

class ChatMessage(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    messages: List[ChatMessage]

# ─────────────────────────────────────────
# HEALTH  (unchanged)
# ─────────────────────────────────────────
@app.get("/")
def home():
    return {"message": "LegalEase AI Backend Running 🚀"}


# ═════════════════════════════════════════
# AUTH ROUTES  (public — no token needed)
# ═════════════════════════════════════════

@app.post("/auth/register")
def do_register(body: RegisterBody):
    result = _auth.register(body.name, body.email, body.password)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result

@app.post("/auth/login")
def do_login(body: LoginBody):
    result = _auth.login(body.email, body.password)
    if "error" in result:
        raise HTTPException(status_code=401, detail=result["error"])
    return result

@app.get("/auth/me")
def get_me(user=Depends(require_user)):
    return {"user": user}

@app.post("/auth/change-password")
def do_change_password(body: ChangePasswordBody, user=Depends(require_user)):
    result = _auth.change_password(user["email"], body.old_password, body.new_password)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result

@app.post("/auth/forgot-password")
def do_forgot_password(body: ForgotPasswordBody):
    """Always returns 200 — never reveals whether email is registered."""
    result    = _auth.create_reset_token(body.email)
    dev_token = result.get("dev_token")
    if dev_token:
        user_rec = _auth._USERS.get(body.email.strip().lower(), {})
        send_reset_email(body.email.strip().lower(), dev_token, user_rec.get("name", ""))
    # Never expose the token in the API response — terminal/email only
    return {
        "success": True,
        "message": "If that email is registered you will receive a reset link shortly.",
    }

@app.post("/auth/reset-password")
def do_reset_password(body: ResetPasswordBody):
    result = _auth.reset_password_with_token(body.token, body.new_password)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result

@app.post("/auth/update-profile")
def do_update_profile(body: UpdateProfileBody, user=Depends(require_user)):
    result = _auth.update_name(user["email"], body.name)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result


# ═════════════════════════════════════════
# PROTECTED ROUTES  (JWT required)
# ═════════════════════════════════════════

# ─────────────────────────────────────────
# FILE TEXT EXTRACTION  (PDF, DOCX, TXT)
# ─────────────────────────────────────────

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}

def _is_allowed(filename: str) -> bool:
    return any(filename.lower().endswith(ext) for ext in ALLOWED_EXTENSIONS)

def extract_pdf_text(file) -> list:
    """Extract text page-by-page from a PDF file."""
    pages_data = []
    with pdfplumber.open(file) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                pages_data.append({"page": i + 1, "text": text})
    return pages_data

def extract_docx_text(file_bytes: bytes) -> list:
    """Extract text paragraph-by-paragraph from a DOCX file."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Group into chunks of ~30 paragraphs to mimic "pages"
    chunk_size = 30
    pages_data = []
    for i in range(0, len(paragraphs), chunk_size):
        chunk = "\n".join(paragraphs[i:i + chunk_size])
        if chunk.strip():
            pages_data.append({"page": i // chunk_size + 1, "text": chunk})
    return pages_data

def extract_txt_text(file_bytes: bytes) -> list:
    """Extract text from a plain text file."""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1", errors="replace")
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    chunk_size = 60
    pages_data = []
    for i in range(0, len(lines), chunk_size):
        chunk = "\n".join(lines[i:i + chunk_size])
        if chunk.strip():
            pages_data.append({"page": i // chunk_size + 1, "text": chunk})
    return pages_data

def extract_text_from_upload(file_bytes: bytes, filename: str) -> list:
    """Route to the correct extractor based on file extension."""
    fname = filename.lower()
    if fname.endswith(".pdf"):
        return extract_pdf_text(io.BytesIO(file_bytes))
    elif fname.endswith(".docx") or fname.endswith(".doc"):
        return extract_docx_text(file_bytes)
    elif fname.endswith(".txt"):
        return extract_txt_text(file_bytes)
    return []

# ─────────────────────────────────────────
# FIND RELEVANT PAGES FOR CHAT  (unchanged)
# ─────────────────────────────────────────
def find_relevant_pages(question, pages, max_pages=3):

    words = question.lower().split()

    scored = []

    for p in pages:
        text = p["text"].lower()

        score = sum(word in text for word in words)

        scored.append((score, p))

    scored.sort(reverse=True, key=lambda x: x[0])

    return [p for score, p in scored[:max_pages] if score > 0]

# ─────────────────────────────────────────
# ANALYZE DOCUMENT  (logic unchanged, JWT guard added)
# ─────────────────────────────────────────
@app.post("/analyze-document")
async def analyze_document(
    file: UploadFile = File(...),
    depth: str = "Standard",
    user=Depends(require_user),
):
    if not _is_allowed(file.filename):
        return {"error": "Only PDF, Word (.docx), and text (.txt) files are supported"}

    # Validate depth value
    if depth not in ("Standard", "Detailed", "Expert"):
        depth = "Standard"

    file_bytes = await file.read()

    try:
        pages = extract_text_from_upload(file_bytes, file.filename)
    except Exception as e:
        print("EXTRACTION ERROR:", e)
        return {"error": f"Could not read the file: {e}"}

    if not pages:
        return {"error": "No readable text found in the document"}

    global DOCUMENT_CONTEXT
    DOCUMENT_CONTEXT = pages

    text = "\n".join([p["text"] for p in pages])

    if not text.strip():
        return {"error": "Document appears to be empty or unreadable"}

    # Fix #10: capture whether text will be truncated before sending to Groq
    DEPTH_LIMITS = {"Expert": 18000, "Detailed": 15000, "Standard": 12000}
    max_chars = DEPTH_LIMITS.get(depth, 12000)
    was_truncated = len(text) > max_chars
    original_char_count = len(text)

    print(f"\n📊 Analysis depth: {depth}")
    ai_response = analyze_legal_text(text, depth=depth)

    # ── Robust cleaning ────────────────────────────────────────────────────
    import re as _re
    clean = ai_response.strip()
    clean = _re.sub(r"```[a-z]*", "", clean).replace("```", "").strip()
    brace = clean.find("{")
    if brace > 0:
        clean = clean[brace:]
    rbrace = clean.rfind("}")
    if rbrace != -1 and rbrace < len(clean) - 1:
        clean = clean[:rbrace + 1]

    parsed = None
    try:
        parsed = json.loads(clean)
    except Exception:
        try:
            fixed = _re.sub(r",\s*}", "}", clean)
            fixed = _re.sub(r",\s*]", "]", fixed)
            parsed = json.loads(fixed)
        except Exception as e:
            print("PARSE FAILED:", e)
            print("RAW:", clean[:500])
            return {"error": "AI response parsing failed — please try again"}

    try:
        summary = parsed.get("simple_summary") or parsed.get("summary", "Analysis complete.")
        if isinstance(summary, list):
            summary = " ".join(summary)
        return {
            "summary":             summary,
            "pros":                parsed.get("pros", []),
            "cons":                parsed.get("cons", []),
            "legal_terms":         parsed.get("legal_terms", []),
            "risk_score":          int(parsed.get("risk_score", 50)),
            "risk_level":          parsed.get("risk_level", "Medium"),
            "suggestions":         parsed.get("suggestions", []),
            "extracted_text":      text,
            "was_truncated":       was_truncated,         # Fix #10
            "original_char_count": original_char_count,  # Fix #10
            "analyzed_char_count": min(original_char_count, max_chars),  # Fix #10
        }
    except Exception as e:
        print("SHAPE ERROR:", e)
        return {"error": "AI response parsing failed — please try again"}


# ─────────────────────────────────────────
# CHAT ENDPOINT  — forwards full message history to Groq
# ─────────────────────────────────────────
@app.post("/chat")
async def chat(req: ChatRequest, user=Depends(require_user)):
    try:
        from groq_client import client as groq_client

        # Build the messages array directly from the request.
        # The frontend sends: [system, ...history, user] so we just forward it.
        messages_payload = [
            {"role": m.role, "content": m.content}
            for m in req.messages
        ]

        # Safety: cap total tokens by trimming oldest non-system messages if payload is very long
        MAX_MESSAGES = 20
        system_msgs  = [m for m in messages_payload if m["role"] == "system"]
        other_msgs   = [m for m in messages_payload if m["role"] != "system"]
        if len(other_msgs) > MAX_MESSAGES:
            other_msgs = other_msgs[-MAX_MESSAGES:]
        messages_payload = system_msgs + other_msgs

        completion = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=messages_payload,
            max_tokens=600,
            temperature=0.4,
        )
        reply = completion.choices[0].message.content.strip()
        return {"reply": reply}
    except Exception as e:
        print("CHAT ERROR:", str(e))
        return {"reply": "I'm having trouble connecting right now. Please try again shortly."}


# ─────────────────────────────────────────
# EXPORT REPORT  (logic unchanged, JWT guard added)
# ─────────────────────────────────────────
@app.post("/export-report")
async def export_report(data: dict, user=Depends(require_user)):

    filename = f"report_{uuid.uuid4().hex}.pdf"

    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=60,
        leftMargin=60,
        topMargin=120,
        bottomMargin=60
    )

    elements = []
    styles = getSampleStyleSheet()

    # ─────────────────────────
    # SAFE KEY HANDLING
    # ─────────────────────────
    risk_score = data.get("risk_score") or data.get("riskScore")
    risk_level = data.get("risk_level") or data.get("riskLabel")

    summary = data.get("summary", "")
    pros = data.get("pros", [])
    cons = data.get("cons", [])
    legal_terms = data.get("legal_terms", [])
    clauses = data.get("clauses", [])  # ⭐ CLAUSE CHECKER

    if isinstance(summary, list):
        summary = " ".join(summary)

    # Convert legal_terms if string format
    if legal_terms and isinstance(legal_terms[0], str):
        formatted = []
        for item in legal_terms:
            if ":" in item:
                term, explanation = item.split(":", 1)
                formatted.append({
                    "section": "Section not clearly defined",
                    "term": term.strip(),
                    "explanation": explanation.strip()
                })
        legal_terms = formatted

    # ─────────────────────────
    # STYLES
    # ─────────────────────────
    heading = ParagraphStyle(
        "Heading",
        parent=styles["Heading2"],
        fontSize=14,
        spaceBefore=15,
        spaceAfter=6
    )

    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontSize=11,
        leading=16
    )

    # ─────────────────────────
    # RISK BADGE
    # ─────────────────────────
    if risk_level == "Low":
        badge_color = "#16a34a"
    elif risk_level == "Medium":
        badge_color = "#f59e0b"
    else:
        badge_color = "#dc2626"

    risk_style = ParagraphStyle(
        "Risk",
        parent=styles["Normal"],
        fontSize=12,
        textColor=colors.white,
        backColor=colors.HexColor(badge_color),
        leftIndent=6,
        rightIndent=6,
        leading=18
    )

    elements.append(
        Paragraph(
            f"<b>Overall Risk Level: {risk_level}</b><br/>Risk Score: {risk_score}/100",
            risk_style
        )
    )

    elements.append(Spacer(1, 0.4 * inch))

    elements.append(Paragraph("Executive Summary", heading))
    elements.append(Paragraph(summary, body))

    elements.append(Paragraph("Favorable Clauses", heading))
    if pros:
        elements.append(
            ListFlowable(
                [ListItem(Paragraph(p, body)) for p in pros],
                bulletType="bullet",
                leftIndent=15
            )
        )

    elements.append(Paragraph("Risks & Red Flags", heading))
    if cons:
        elements.append(
            ListFlowable(
                [ListItem(Paragraph(c, body)) for c in cons],
                bulletType="bullet",
                leftIndent=15
            )
        )

    elements.append(Paragraph("Legal Terms Explained", heading))

    for term in legal_terms:

        section = term.get("section", "Section not clearly defined")
        term_name = term.get("term", "Legal Term")
        explanation = term.get("explanation", "")

        elements.append(
            Paragraph(
                f"<b>{section} — {term_name}</b>",
                styles["Heading4"]
            )
        )
        elements.append(Paragraph(explanation, body))
        elements.append(Spacer(1, 0.2 * inch))

    elements.append(Paragraph("Clause-by-Clause Report", heading))

    for clause in clauses:

        title = clause.get("title", "")
        note = clause.get("note", "")
        status = clause.get("status", "ok")

        if status == "ok":
            label = "Fair"
        elif status == "warn":
            label = "Review"
        else:
            label = "Risk"

        elements.append(Paragraph(f"<b>[{label}] {title}</b>", body))
        elements.append(Paragraph(note, body))
        elements.append(Spacer(1, 0.2 * inch))

    def add_layout(canvas, doc):
        canvas.saveState()

        canvas.setFillColor(colors.HexColor("#1e3a8a"))
        canvas.rect(0, A4[1] - 70, A4[0], 70, fill=1)

        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 16)
        canvas.drawString(60, A4[1] - 45, "LegalEase AI")

        canvas.setFont("Helvetica", 80)
        canvas.setFillColorRGB(0.95, 0.95, 0.95)
        canvas.drawCentredString(A4[0] / 2, A4[1] / 2, "LegalEase")

        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(colors.grey)
        canvas.drawRightString(A4[0] - 60, 30, f"Page {doc.page}")

        canvas.restoreState()

    doc.build(elements, onFirstPage=add_layout, onLaterPages=add_layout)

    return FileResponse(
        filename,
        media_type="application/pdf",
        filename="LegalEase_Report.pdf"
    )


# ─────────────────────────────────────────
# EXPORT DRAFTED DOCUMENT AS PDF
# ─────────────────────────────────────────
class ExportDocBody(BaseModel):
    title: str
    content: str

@app.post("/export-doc-pdf")
async def export_doc_pdf(body: ExportDocBody, user=Depends(require_user)):
    filename = f"doc_{uuid.uuid4().hex}.pdf"
    doc = SimpleDocTemplate(filename, pagesize=A4,
                            rightMargin=60, leftMargin=60,
                            topMargin=80, bottomMargin=60)
    styles = getSampleStyleSheet()
    elements = []

    title_style = ParagraphStyle("DocTitle", parent=styles["Title"],
                                 fontSize=16, spaceAfter=12,
                                 textColor=colors.HexColor("#1e3a8a"))
    body_style  = ParagraphStyle("DocBody",  parent=styles["BodyText"],
                                 fontSize=11, leading=18, spaceAfter=4)
    section_style = ParagraphStyle("Sec", parent=styles["Heading3"],
                                   fontSize=12, spaceAfter=4, spaceBefore=10,
                                   textColor=colors.HexColor("#1e3a8a"))

    elements.append(Paragraph(body.title, title_style))
    elements.append(Spacer(1, 0.2 * inch))

    for line in body.content.split("\n"):
        stripped = line.strip()
        if not stripped:
            elements.append(Spacer(1, 0.1 * inch))
        elif stripped.isupper() and len(stripped) < 80:
            elements.append(Paragraph(f"<b>{stripped}</b>", section_style))
        else:
            safe = stripped.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            elements.append(Paragraph(safe, body_style))

    def header_footer(canvas, doc_obj):
        canvas.saveState()

        # ── Header bar ──
        header_h = 50
        canvas.setFillColor(colors.HexColor("#1e3a8a"))
        canvas.rect(0, A4[1] - header_h, A4[0], header_h, fill=1, stroke=0)

        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 15)
        canvas.drawString(48, A4[1] - 32, "LegalEase AI")

        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(colors.HexColor("#a0b4e0"))
        canvas.drawString(48, A4[1] - 44, "Smart Property Document Platform")

        # Page number on the right side of header
        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica", 9)
        canvas.drawRightString(A4[0] - 48, A4[1] - 32, body.title[:55])
        canvas.drawRightString(A4[0] - 48, A4[1] - 44, f"Page {doc_obj.page}")

        # ── Thin accent line below header ──
        canvas.setStrokeColor(colors.HexColor("#3b5fc0"))
        canvas.setLineWidth(1)
        canvas.line(0, A4[1] - header_h - 1, A4[0], A4[1] - header_h - 1)

        # ── Footer ──
        canvas.setFillColor(colors.HexColor("#888888"))
        canvas.setFont("Helvetica", 8)
        canvas.drawCentredString(A4[0] / 2, 22,
            "Generated by LegalEase AI  |  For informational purposes only  |  Consult a legal professional before signing")

        canvas.restoreState()

    doc.build(elements, onFirstPage=header_footer, onLaterPages=header_footer)
    return FileResponse(filename, media_type="application/pdf",
                        filename=f"{body.title}.pdf")


# ─────────────────────────────────────────
# EXPORT DRAFTED DOCUMENT AS DOCX
# ─────────────────────────────────────────
@app.post("/export-doc-docx")
async def export_doc_docx(body: ExportDocBody, user=Depends(require_user)):
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    filename = f"doc_{uuid.uuid4().hex}.docx"
    document = DocxDocument()

    # Page margins
    for section in document.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Title
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(body.title)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)

    document.add_paragraph()

    for line in body.content.split("\n"):
        stripped = line.strip()
        if not stripped:
            document.add_paragraph()
        elif stripped.isupper() and len(stripped) < 80:
            p = document.add_paragraph()
            run = p.add_run(stripped)
            run.bold      = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
        else:
            p = document.add_paragraph(stripped)
            p.runs[0].font.size = Pt(11) if p.runs else None

    document.save(filename)
    return FileResponse(filename,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        filename=f"{body.title}.docx")


# ─────────────────────────────────────────
# ANALYSE DRAFTED DOCUMENT TEXT (via Groq)
# ─────────────────────────────────────────
class AnalyseTextBody(BaseModel):
    content: str

@app.post("/analyse-draft")
async def analyse_draft(body: AnalyseTextBody, user=Depends(require_user)):
    import re as _re
    ai_response = analyze_legal_text(body.content)
    clean = ai_response.strip()
    clean = _re.sub(r"```[a-z]*", "", clean).replace("```", "").strip()
    brace = clean.find("{")
    if brace > 0:
        clean = clean[brace:]
    rbrace = clean.rfind("}")
    if rbrace != -1 and rbrace < len(clean) - 1:
        clean = clean[:rbrace + 1]
    try:
        parsed = json.loads(clean)
    except Exception:
        try:
            fixed = _re.sub(r",\s*}", "}", clean)
            fixed = _re.sub(r",\s*]", "]", fixed)
            parsed = json.loads(fixed)
        except Exception:
            return {"error": "AI parsing failed — please try again"}
    # Detect groq error marker
    if parsed.get("risk_score") == -1 or "GROQ_ERROR" in str(parsed.get("simple_summary","")):
        err_detail = str(parsed.get("simple_summary","")).replace("GROQ_ERROR: ","")
        return {"error": f"AI service error: {err_detail}"}
    summary = parsed.get("simple_summary") or parsed.get("summary", "Analysis complete.")
    if isinstance(summary, list):
        summary = " ".join(summary)
    return {
        "summary":    summary,
        "pros":       parsed.get("pros", []),
        "cons":       parsed.get("cons", []),
        "legal_terms":parsed.get("legal_terms", []),
        "risk_score": int(parsed.get("risk_score", 50)),
        "risk_level": parsed.get("risk_level", "Medium"),
        "suggestions":parsed.get("suggestions", []),
    }


# ─────────────────────────────────────────
# SEND E-SIGN REQUEST EMAIL
# ─────────────────────────────────────────
class ESignBody(BaseModel):
    to_email:     str
    to_name:      str
    from_name:    str
    doc_title:    str
    message:      str
    doc_content:  str

@app.post("/send-sign-request")
async def send_sign_request(body: ESignBody, user=Depends(require_user)):
    from email.mime.base import MIMEBase
    from email import encoders
    import tempfile

    # Build a simple PDF attachment of the document
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp_name = tmp.name
    tmp.close()

    doc_pdf = SimpleDocTemplate(tmp_name, pagesize=A4,
                                rightMargin=72, leftMargin=72,
                                topMargin=100, bottomMargin=60)
    styles   = getSampleStyleSheet()
    elements = []
    body_style = ParagraphStyle("B", parent=styles["BodyText"], fontSize=11, leading=18)
    elements.append(Paragraph(f"<b>{body.doc_title}</b>",
        ParagraphStyle("T", parent=styles["Title"], fontSize=15,
                       textColor=colors.HexColor("#1e3a8a"), spaceAfter=14)))
    elements.append(Spacer(1, 0.15 * inch))
    for line in body.doc_content.split("\n"):
        stripped = line.strip()
        if not stripped:
            elements.append(Spacer(1, 0.08 * inch))
        elif stripped.isupper() and len(stripped) < 80:
            elements.append(Paragraph(f"<b>{stripped.replace('&','&amp;').replace('<','&lt;')}</b>",
                ParagraphStyle("S", parent=styles["Heading3"], fontSize=11,
                               textColor=colors.HexColor("#1e3a8a"), spaceAfter=4)))
        else:
            safe = stripped.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            elements.append(Paragraph(safe, body_style))
    doc_pdf.build(elements)

    # Read PDF bytes
    with open(tmp_name, "rb") as f:
        pdf_bytes = f.read()

    greeting  = f"Dear {body.to_name},"
    user_msg  = body.message or f"Please review and sign the attached document at your earliest convenience."
    sign_note = "This document has been shared with you via LegalEase AI for your review and signature. Please sign and return a copy to the sender."

    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
<style>
body{{font-family:'Segoe UI',Arial,sans-serif;background:#0d1117;margin:0;padding:24px}}
.card{{max-width:520px;margin:0 auto;background:#161b22;border-radius:16px;border:1px solid rgba(48,54,70,.9);overflow:hidden}}
.top{{background:linear-gradient(135deg,#4f6ef7,#6d28d9);padding:26px 30px}}
.logo{{font-size:17px;font-weight:800;color:#fff}}
.sub{{font-size:12px;color:rgba(255,255,255,.7);margin-top:4px}}
.body{{padding:26px 30px}}
h2{{font-size:18px;font-weight:700;color:#e6edf3;margin:0 0 12px}}
p{{font-size:13px;color:#8b949e;line-height:1.7;margin:0 0 12px}}
.doc-card{{background:#0d1117;border:1px solid rgba(48,54,70,.9);border-radius:10px;padding:14px 16px;margin:16px 0}}
.doc-name{{font-size:14px;font-weight:700;color:#e6edf3}}
.doc-sub{{font-size:11px;color:#484f58;margin-top:3px}}
.note{{background:#1a2e1a;border:1px solid #1a5c32;border-radius:8px;padding:12px 14px;font-size:12px;color:#4ade80;margin-top:16px}}
.foot{{font-size:11px;color:#484f58;border-top:1px solid rgba(48,54,70,.9);padding:16px 30px;background:#0d1117}}
</style></head><body>
<div class="card">
  <div class="top">
    <div class="logo">⚖️ LegalEase AI</div>
    <div class="sub">Document Signing Request</div>
  </div>
  <div class="body">
    <h2>You have a document to sign</h2>
    <p>{greeting}</p>
    <p><strong style="color:#e6edf3">{body.from_name}</strong> has shared a document with you and is requesting your signature.</p>
    <p style="font-style:italic;color:#6b7280">{user_msg}</p>
    <div class="doc-card">
      <div class="doc-name">📄 {body.doc_title}</div>
      <div class="doc-sub">Attached as PDF · Prepared via LegalEase AI</div>
    </div>
    <div class="note">✅ {sign_note}</div>
    <p style="margin-top:16px;font-size:12px">Please sign and return this document to <strong style="color:#e6edf3">{user.get('email','')}</strong>.</p>
  </div>
  <div class="foot">LegalEase AI · This is an automated signing request. Not a substitute for qualified legal advice.</div>
</div></body></html>"""

    plain = (
        f"{greeting}\n\n{body.from_name} has sent you '{body.doc_title}' for signing.\n\n"
        f"{user_msg}\n\nThe document is attached as a PDF.\n\n"
        f"Please sign and return to {user.get('email','')}.\n\nSent via LegalEase AI."
    )

    # ── Send via Brevo SMTP relay ─────────────────────────────────────────────
    safe_title  = body.doc_title.replace(" ", "_")[:50]
    attachments = [(f"{safe_title}.pdf", pdf_bytes)]  # (filename, bytes) tuples

    try:
        sent = send_email(
            to          = body.to_email,
            subject     = f"Signature Requested: {body.doc_title}",
            html        = html,
            plain       = plain,
            attachments = attachments,
        )
        if not sent:
            print(f"\n[DEV MODE] E-sign would be sent to {body.to_email} — set BREVO_SMTP_USER + BREVO_SMTP_PASS on Railway")
            return {"success": True, "dev": True, "message": "Dev mode — set BREVO_SMTP_USER and BREVO_SMTP_PASS to send real emails"}
        return {"success": True, "message": f"Signing request sent to {body.to_email}"}
    except Exception as e:
        print(f"[EMAIL ERROR] {type(e).__name__}: {e}")
        raise HTTPException(status_code=500, detail=f"Email send failed: {str(e)}")

# ═══════════════════════════════════════════════════════════════════════════════
# HISTORY ENDPOINTS  (Supabase-backed — replaces localStorage on frontend)
# ═══════════════════════════════════════════════════════════════════════════════

# ── Pydantic models ────────────────────────────────────────────────────────────
class AnalysisHistoryEntry(BaseModel):
    id:         Optional[str] = None
    name:       str
    size:       Optional[str] = None
    risk:       Optional[str] = None
    risk_score: Optional[int] = None
    depth:      Optional[str] = None
    date:       Optional[str] = None
    starred:    bool = False
    result:     Optional[dict] = None

class DraftHistoryEntry(BaseModel):
    id:          Optional[str] = None
    name:        str
    category:    Optional[str] = None
    content:     Optional[str] = None
    sign_status: Optional[str] = "Draft"
    ai_score:    Optional[int] = None
    starred:     bool = False
    date:        Optional[str] = None

class StarBody(BaseModel):
    starred: bool

class DeleteAllBody(BaseModel):
    type: str  # "analysis" | "drafts" | "all"


def _require_supabase():
    if not _sb:
        raise HTTPException(status_code=503, detail="Database not configured. Add SUPABASE_URL and SUPABASE_SERVICE_KEY to .env.local")


# ── Analysis History ───────────────────────────────────────────────────────────

@app.get("/history/analysis")
def get_analysis_history(user=Depends(require_user)):
    _require_supabase()
    try:
        res = _sb.table("analysis_history")\
            .select("*")\
            .eq("user_email", user["email"])\
            .order("created_at", desc=True)\
            .execute()
        return {"entries": res.data or []}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/history/analysis")
def save_analysis_entry(entry: AnalysisHistoryEntry, user=Depends(require_user)):
    _require_supabase()
    try:
        row = {
            "id":         entry.id or str(uuid.uuid4()),
            "user_email": user["email"],
            "name":       entry.name,
            "size":       entry.size,
            "risk":       entry.risk,
            "risk_score": entry.risk_score,
            "depth":      entry.depth,
            "date":       entry.date,
            "starred":    entry.starred,
            "result":     entry.result,
        }
        res = _sb.table("analysis_history").upsert(row).execute()
        return {"entry": res.data[0] if res.data else row}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.patch("/history/analysis/{entry_id}/star")
def star_analysis_entry(entry_id: str, body: StarBody, user=Depends(require_user)):
    _require_supabase()
    try:
        _sb.table("analysis_history")\
            .update({"starred": body.starred})\
            .eq("id", entry_id)\
            .eq("user_email", user["email"])\
            .execute()
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/history/analysis/{entry_id}")
def delete_analysis_entry(entry_id: str, user=Depends(require_user)):
    _require_supabase()
    try:
        _sb.table("analysis_history")\
            .delete()\
            .eq("id", entry_id)\
            .eq("user_email", user["email"])\
            .execute()
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── Draft History ──────────────────────────────────────────────────────────────

@app.get("/history/drafts")
def get_draft_history(user=Depends(require_user)):
    _require_supabase()
    try:
        res = _sb.table("draft_history")\
            .select("*")\
            .eq("user_email", user["email"])\
            .order("created_at", desc=True)\
            .execute()
        return {"entries": res.data or []}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/history/drafts")
def save_draft_entry(entry: DraftHistoryEntry, user=Depends(require_user)):
    _require_supabase()
    try:
        row = {
            "id":          entry.id or str(uuid.uuid4()),
            "user_email":  user["email"],
            "name":        entry.name,
            "category":    entry.category,
            "content":     entry.content,
            "sign_status": entry.sign_status or "Draft",
            "ai_score":    entry.ai_score,
            "starred":     entry.starred,
            "date":        entry.date,
        }
        res = _sb.table("draft_history").upsert(row).execute()
        return {"entry": res.data[0] if res.data else row}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.patch("/history/drafts/{entry_id}/star")
def star_draft_entry(entry_id: str, body: StarBody, user=Depends(require_user)):
    _require_supabase()
    try:
        _sb.table("draft_history")\
            .update({"starred": body.starred})\
            .eq("id", entry_id)\
            .eq("user_email", user["email"])\
            .execute()
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/history/drafts/{entry_id}")
def delete_draft_entry(entry_id: str, user=Depends(require_user)):
    _require_supabase()
    try:
        _sb.table("draft_history")\
            .delete()\
            .eq("id", entry_id)\
            .eq("user_email", user["email"])\
            .execute()
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── Clear all history ──────────────────────────────────────────────────────────

@app.delete("/history/all")
def clear_all_history(user=Depends(require_user)):
    _require_supabase()
    try:
        _sb.table("analysis_history").delete().eq("user_email", user["email"]).execute()
        _sb.table("draft_history").delete().eq("user_email", user["email"]).execute()
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
